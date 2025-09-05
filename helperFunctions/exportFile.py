# generate an excel sheet from data
from fastapi import  Response
from openpyxl import Workbook
from io import BytesIO
from typing import Any
import os
from docx import Document
from openpyxl.styles import Font
import re

# import inflection

# def prettify_key(text):
#     return inflection.titleize(inflection.underscore(text))

    
def prettify_key(key):
    if key == "memberID":
        return key  # Keep it exactly as it is
    
    # Insert space before each uppercase letter except if it's the first character
    words = re.sub(r'(?<!^)(?=[A-Z])', ' ', key)
    # Capitalize the first letter of each word
    title_case = words.title()
    
    return title_case


def generate_excel(data: Any, filename : str):
    """
    Generate an Excel file from any data.
    :param data: Data of any type (list, dictionary, etc.).
    :return: Excel file as a downloadable response.
    """
    # Create an Excel workbook and worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = filename

    print("work book created ", wb)
    print("work book is active ", ws)
    print("work book name ", ws.title)

    # Determine the structure of the data and process it
    if isinstance(data, list):
        print("am in the list ")
        if not data:
            raise HTTPException(status_code=400, detail="The list is empty.")
        if isinstance(data[0], dict):
            # List of dictionaries
            headers = list(data[0].keys())
            print ("headers :", headers)
            pretty_headers = [prettify_key(key) for key in headers]
            print ("p headers :", pretty_headers)

            for col_num, header in enumerate(pretty_headers, start=1):
                cell = ws.cell(row=1, column=col_num, value=header)
                cell.font = Font(bold=True)


            for row in data:
                ws.append(list(row.values()))
        else:
            # List of non-dictionary elements
            headers = ["Values"]
            ws.append(headers)
            for item in data:
                ws.append([item])

    elif isinstance(data, dict):
        print("am in the dict ")
        # Single dictionary
        headers = list(data.keys())
        pretty_headers = [prettify_key(key) for key in headers]
        ws.append(pretty_headers)
        ws.append(list(data.values()))

    else:
        print("am in the other type ", type(data))
        # Fallback for other types
        headers = ["Data"]
        ws.append(headers)
        ws.append([data])

    print("excel data ",wb)
    
    # Use the existing 'uploads' folder
    upload_folder_path = os.path.join(os.path.dirname(__file__), 'reports')

    # Ensure the 'uploads' directory exists (optional, in case it wasn't created)
    os.makedirs(upload_folder_path, exist_ok=True)

    # Define the path to save the file in your existing 'uploads' folder
    save_path = os.path.join(upload_folder_path, f"{filename}.xlsx")
    # Create a temporary file to save the workbook
    # temp_file_path = f"/tmp/{filename}.xlsx"  # You can change the path as needed
    wb.save(save_path)

    # output = BytesIO()
    # wb.save(output)
    # output.seek(0)

    # Create a response with the Excel file
    # response_headers = {
    #     'Content-Disposition': 'attachment; filename="data.xlsx"'
    # }
    # return Response(
    #     content=output.read(),
    #     media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    #     headers=response_headers
    # )

    return save_path


    # def generate_docx(data: Any, filename : str):
    #     document = Document()

    #     document.add_heading('CTC Member Data', level=1)
    data = {
        'title': '', 'firstName': 'Millicent', 'middleName': '', 'lastName': 'Ocloo', 'dateOfBirth': '',
        'age': '', 'gender': 'Female', 'phoneNumber': '0592634014', 'email': '', 'nationality': '',
        'homeTown': '', 'homeAddress': 'Kordiabe', 'workingStatus': '', 'occupation': '', 'qualification': '',
        'institutionName': '', 'mothersName': '', 'fathersName': '', 'nextOfKin': '', 'nextOfKinPhoneNumber': '',
        'maritalStatus': '', 'spouseContact': '', 'spouseName': '', 'numberOfChildren': '', 'memberType': '',
        'cell': '', 'departmentName': '', 'dateJoined': '', 'classSelection': '', 'spiritualGift': '', 'position': '',
        'waterBaptised': '', 'baptisedBy': '', 'dateBaptised': '', 'baptisedByTheHolySpirit': '', 'memberStatus': '',
        'dateDeceased': '', 'dateBuried': '', 'confirmed': '', 'dateConfirmed': '', 'comment': ''
    }

# def prettify_key(key):
#     return key.replace("_", " ").replace("-", " ").title().replace("Of", "of").replace("ByThe", "by the")

def generate_docx():

        data = {
            'title': '', 'firstName': 'Millicent', 'middleName': '', 'lastName': 'Ocloo', 'dateOfBirth': '',
            'age': '', 'gender': 'Female', 'phoneNumber': '0592634014', 'email': '', 'nationality': '',
            'homeTown': '', 'homeAddress': 'Kordiabe', 'workingStatus': '', 'occupation': '', 'qualification': '',
            'institutionName': '', 'mothersName': '', 'fathersName': '', 'nextOfKin': '', 'nextOfKinPhoneNumber': '',
            'maritalStatus': '', 'spouseContact': '', 'spouseName': '', 'numberOfChildren': '', 'memberType': '',
            'cell': '', 'departmentName': '', 'dateJoined': '', 'classSelection': '', 'spiritualGift': '', 'position': '',
            'waterBaptised': '', 'baptisedBy': '', 'dateBaptised': '', 'baptisedByTheHolySpirit': '', 'memberStatus': '',
            'dateDeceased': '', 'dateBuried': '', 'confirmed': '', 'dateConfirmed': '', 'comment': ''
        }


        # Create a new Word document
        doc = Document()
        doc.add_heading('Member Information Form', level=1)

        # Add each field in its own row in a 2-column table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        number = 1
        for key, value in data.items():
            row = table.add_row().cells
            row[0].text = f"{number}. {prettify_key(key)}"
            row[1].text = value if value else "-"
            number += 1

        # Save the document
        upload_folder_path = os.path.join(os.path.dirname(__file__), 'reports')
        print("file location : ",upload_folder_path ) 
        saved_path = os.path.join(upload_folder_path, "Formatted_test_Info.docx")
        doc.save(saved_path)
        print("saved path : ", saved_path)
        print("Formatted document saved as 'Formatted_test_Info.docx'")

        return saved_path