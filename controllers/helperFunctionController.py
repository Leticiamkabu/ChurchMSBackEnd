from fastapi import APIRouter, Depends, HTTPException, FastAPI, HTTPException 
from typing import Annotated
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy import select, or_, func, and_, create_engine, MetaData, Table, update, null, insert , inspect
from sqlalchemy.sql import func , extract
from sqlalchemy.engine.result import Row
from datetime import datetime
from dotenv import load_dotenv

import logging
import os
import subprocess
import tempfile
from fastapi import UploadFile,status, File, Form, Request
import uuid


# create a connection to the database
async def get_db():
    db = SessionLocal()
    try:
        yield db

    finally:
        await db.close()

db_dependency = Annotated[Session, Depends(get_db)]




router = APIRouter()

# Load environment variables from .env file
load_dotenv()

# Fetch database credentials from environment variables
db_user = os.getenv('DB_USER')
db_pass = os.getenv('DB_PASS')
db_host = os.getenv('DB_HOST')
db_name = os.getenv('DB_NAME')




@router.post("/helperFunction/clean_data_base/{table_name}")
async def replace_null_values_in_the_database(table_name: str):

    DATABASE_URL = "postgresql://ctc_dev_pejs_user:tQ8MR9iUIu4hJeeXVMt9lCb8pujCWpZ2@dpg-d6g5rv3h46gs738pbf5g-a.oregon-postgres.render.com/ctc_dev_pejs"

    # DATABASE_URL = f"postgresql://{db_user}:{db_pass}@{db_host}/{db_name}"
    engine = create_engine(DATABASE_URL)
    SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

    session = SessionLocal()
    metadata = MetaData()

    try:
        # Reflect the table
        table = Table(table_name, metadata, autoload_with=engine)

        # Select all rows
        stmt = select(table)
        results = session.execute(stmt).fetchall()

        updated_rows = 0

        for row in results:
            update_data = {}
            colum_numbers = 0
            for column in table.columns:

                
                if row[colum_numbers] is None:
                    update_data[column.name] = ""
                colum_numbers += 1

        
            if update_data:
                upd = (
                    update(table)
                    .where(table.columns.id == row[0])
                    .values(**update_data)
                )
                session.execute(upd)
                updated_rows += 1

        session.commit()
        return {"message": f"{updated_rows} row(s) updated."}

    except Exception as e:
        session.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        session.close()


# work on this soon
# postgresql://ctc_dev_pejs_user:tQ8MR9iUIu4hJeeXVMt9lCb8pujCWpZ2@dpg-d6g5rv3h46gs738pbf5g-a.oregon-postgres.render.com/ctc_dev_pejs
@router.post("/helperFunction/clone_db")
async def clone_db():
    try:
        # Dump source DB to temp file
        with tempfile.NamedTemporaryFile(delete=False) as tmpfile:
            dump_file = tmpfile.name

        # Source DB credentials
        env = os.environ.copy()
        env["PGPASSWORD"] =  "leeminho"
        env["PGSSLMODE"] = "disable"

        dump_cmd = [
            "pg_dump",
            "-h",  "localhost",
            "-p", str(5432),
            "-U",  "postgres",
            "-d",  "ctcTest",
            "-Fc",
            "--no-owner",
            "-f", dump_file
        ]
        subprocess.run(dump_cmd, check=True, env=env)

        # Target DB credentials
        env["PGPASSWORD"] =  "tQ8MR9iUIu4hJeeXVMt9lCb8pujCWpZ2"
        env["PGSSLMODE"] = "require"

        # restore_cmd = [
        #     "pg_restore",
        #     "-h", "dpg-d4s17dh5pdvs73btcoi0-a.oregon-postgres.render.com",
        #     "-p", str(5432),
        #     "-U", "ctc_dev_n0qw_user",
        #     "-d", "ctc_dev_n0qw",
        #     "--no-owner",
        #     "--no-acl", 
        #     "--clean",
        #     "--verbose",
        #     dump_file
        # ]

        restore_cmd = [
            "pg_restore",
            "-h", "dpg-d6g5rv3h46gs738pbf5g-a.oregon-postgres.render.com",
            "-p", str(5432),
            "-U", "ctc_dev_pejs_user",
            "-d", "ctc_dev_pejs",
            "--no-owner",
            "--no-acl", 
            "--clean",
            "--verbose",
            dump_file
        ]
        subprocess.run(restore_cmd, check=True, env=env)

        os.remove(dump_file)

        return {"status": "success", "message": "Database cloned succesfully"}

    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"Command failed: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@router.post("/helperFunction/move_table_data")
async def data_transfere(table_name: str):   
    
        
    DESTINATION_DB_URL = "postgresql://leticia:leeminho@localhost/ctcTest"
    SOURCE_DB_URL = "postgresql://leticia:leeminho@localhost/churchMSDev"

    # Replace 'your_table_name' with the actual table name you want to transfer
    TABLE_NAME = table_name

    # Set up the engines and metadata
    source_engine = create_engine(SOURCE_DB_URL)
    destination_engine = create_engine(DESTINATION_DB_URL)
    metadata = MetaData()

    # Reflect the table from the source database
    source_table = Table(TABLE_NAME, metadata, autoload_with=source_engine)

    # Connect to both databases
    with source_engine.connect() as source_conn, destination_engine.connect() as dest_conn:
        # Start a transaction in the destination database
        with dest_conn.begin():
             # Check if the destination table exists using the inspector
            inspector = inspect(destination_engine)
            if TABLE_NAME in inspector.get_table_names():
                print(f"Table '{TABLE_NAME}' exists in the destination database.")
                destination_table = Table(TABLE_NAME, metadata, autoload_with=destination_engine)
            else:
                print(f"Table '{TABLE_NAME}' does not exist. Creating table...")
                # Define the schema based on the source table
                destination_table = Table(
                    TABLE_NAME,
                    metadata,
                    *(Column(col.name, col.type, primary_key=col.primary_key) for col in source_table.columns),
                    extend_existing=True

                )
                metadata.create_all(destination_engine)

            # Load data from the source table
            print("Fetching data from source table...")
            select_st = select(source_table)
            results = source_conn.execute(select_st).fetchall()

            # Insert data into the destination table
            print(f"Transferring {len(results)} rows to destination table...")
            for row in results:
                # Convert Row object to dictionary
                row_dict = dict(row._asdict())

                # Insert data into the destination table
                insert_st = insert(destination_table).values(**row_dict)
                dest_conn.execute(insert_st)

    print("Data transfer complete.")
    return "Data transfer complete."

import aiofiles
from docx import Document
import io


def read_docx(doc):
    """Extracts name and department from a Word Document table."""
    data_list = []  # list of dicts for multiple rows

    # Check if the document has tables
    if not doc.tables:
        return {"error": "No tables found in document"}

    table = doc.tables[0]  # assuming first table contains your data

    # Get headers from first row
    headers = [cell.text.strip().upper() for cell in table.rows[0].cells]

    # Check required columns
    if "NAME" not in headers or "DEPARTMENT" not in headers:
        return {"error": "Table must contain 'NAME' and 'DEPARTMENT' columns"}

    name_idx = headers.index("NAME")
    dept_idx = headers.index("DEPARTMENT")

    # Read remaining rows
    for row in table.rows[1:]:
        name = row.cells[name_idx].text.strip()
        department = row.cells[dept_idx].text.strip()
        data_list.append({"NAME": name, "DEPARTMENT": department})

    return data_list


import psycopg2
from psycopg2 import OperationalError
# postgresql://ctc_dev_n0qw_user:sOg4Ja1CNTTdlVr8j15HplPQJQyM94kz@dpg-d4s17dh5pdvs73btcoi0-a.oregon-postgres.render.com/ctc_dev_n0qw
def get_connection():
    """Safely connect to the Render PostgreSQL database with SSL enabled."""
    SOURCE_DB_URL = (
        "postgresql://leticia:"
        "leeminho@"
        "localhost:5432/"
        "ctcTest?"
    )
    try:
        conn = psycopg2.connect(SOURCE_DB_URL)
        return conn
    except OperationalError as e:
        print("⚠️ Database connection failed:", e)
        return None

# sslmode=require
def update_department_in_db(name, department):
    """Updates department for a given employee name in the database."""
    conn = get_connection()
    if not conn:
        return {"status": "error", "message": "Failed to connect to database"}

    cursor = conn.cursor()

    try:
        # Normalize name (lowercase + remove spaces)
        clean_name = name.replace(" ", "").lower().strip()

        # Find matching members
        cursor.execute("""
            SELECT id, "firstName", "middleName", "lastName"
            FROM members
            WHERE LOWER(TRIM("lastName" || COALESCE("middleName", '') || "firstName")) = %s
        """, (clean_name,))

        records = cursor.fetchall()

        if len(records) == 0:
            msg = f"No record found for {name}"
            return {"status": "none", "message": msg}

        elif len(records) > 1:
            msg = f"Multiple records found for {name}, skipping update."
            return {"status": "multiple", "message": msg}

        # Exactly one match → update department
        member_id = records[0][0]
        cursor.execute(
            'UPDATE members SET "departmentName" = %s WHERE id = %s',
            (department, member_id),
        )
        conn.commit()

        msg = f"Department updated for {name}"
        return {"status": "updated", "message": msg}

    except OperationalError as e:
        msg = f"Database operation failed for {name}: {e}"
        print("⚠️", msg)
        return {"status": "error", "message": msg}

    finally:
        cursor.close()
        conn.close()

def generatedId(lastNumber):
    prefix = "CTCAG"
    padded_number = "00"+ str(lastNumber + 1)  # pads to 4 digits like 0001
    year_suffix = str(datetime.now().year)[-2:]  # gets last 2 digits of current year
    return f"{prefix}/{padded_number}/{year_suffix}"

def add_member_in_db(firstName,middleName, lastName, gender,phoneNumber, monthBorn, age, occupation, departmentName, memberStatus):
    """Updates department for a given employee name in the database."""
    conn = get_connection()
    if not conn:
        return {"status": "error", "message": "Failed to connect to database"}

    cursor = conn.cursor()

    try:
        # Normalize name (lowercase + remove spaces)
        # clean_name = name.replace(" ", "").lower().strip()

        # Find matching members
        # cursor.execute("""
        #     SELECT id, "firstName", "middleName", "lastName"
        #     FROM members
        #     WHERE LOWER(TRIM("lastName" || COALESCE("middleName", '') || "firstName")) = %s
        # """, (clean_name,))

        # records = cursor.fetchall()

        # if len(records) == 0:
        #     msg = f"No record found for {name}"
        #     return {"status": "none", "message": msg}

        # elif len(records) > 1:
        #     msg = f"Multiple records found for {name}, skipping update."
        #     return {"status": "multiple", "message": msg}

        # Exactly one match → update department
        cursor.execute("SELECT COUNT(id) FROM members")
        count = cursor.fetchone()[0]

        member_id = generatedId(count)
        cursor.execute(
            """INSERT INTO members ("memberID","firstName", "middleName", "lastName", "gender","phoneNumber", "monthBorn", "age", "occupation", "departmentName", "memberStatus" ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
            (member_id, firstName, middleName, lastName, gender, phoneNumber,monthBorn, age, occupation, departmentName,memberStatus )
            )

        conn.commit()

        msg = f"Member updated for {firstName}"
        return {"status": "added", "message": msg}

    except OperationalError as e:
        msg = f"Database operation failed for {name}: {e}"
        print("⚠️", msg)
        return {"status": "error", "message": msg}

    finally:
        cursor.close()
        conn.close()

import time

@router.post("/helperFunction/updateDepartment")
async def department_data_update(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        return {"error": "Invalid file type. Please upload a .docx file"}

    contents = await file.read()
    doc = Document(io.BytesIO(contents))
    data = read_docx(doc)
    print("Extracted data:", data)

    if "NAME" not in data[0] or "DEPARTMENT" not in data[0]:
        return {"error": "Document must contain 'Name:' and 'Department:' fields"}

    # result = update_department_in_db(data["NAME"], data["DEPARTMENT"])

    for row in data:
        print("rowss : ", row )
        name = row.get("NAME")
        department = row.get("DEPARTMENT")

        if not name or not department:
            results.append({"name": name, "status": "error", "message": "Missing name or department"})
            continue

        update_result = update_department_in_db(name, department)
        time.sleep(20) 

    # Return the appropriate message based on status
        if update_result["status"] == "updated":
            print("message:", update_result["message"], "department:", row.get("DEPARTMENT"))
        else:
            print("error:", update_result["message"], "department:", row.get("DEPARTMENT"))



def read_docx1(doc):
    """Extracts name and department from a Word Document table."""
    data_list = []  # list of dicts for multiple rows

    # Check if the document has tables
    if not doc.tables:
        return {"error": "No tables found in document"}

    table = doc.tables[0]  # assuming first table contains your data

    # Get headers from first row
    headers = [cell.text.strip().upper() for cell in table.rows[0].cells]
    print(headers)

    # Check required columns
    # if "NAME" not in headers or "DEPARTMENT" not in headers:
    #     return {"error": "Table must contain 'NAME' and 'DEPARTMENT' columns"}

    first_name_idx = headers.index("FIRSTNAME")
    middle_name_idx = headers.index("MIDDLENAME")
    last_name_idx = headers.index("LASTNAME")
    gender_idx = headers.index("GENDER")
    phone_number_idx = headers.index("PHONENUMBER")
    month_born_idx = headers.index("MONTHBORN")
    age_idx = headers.index("AGE")
    occupation_idx = headers.index("OCCUPATION")
    department_name_idx = headers.index("DEPARTMENTNAME")
    member_status_idx = headers.index("MEMBERSTATUS")

    # Read remaining rows
    for row in table.rows[1:]:
        first_name  = row.cells[first_name_idx].text.strip()
        middle_name = row.cells[middle_name_idx].text.strip()
        last_name = row.cells[last_name_idx].text.strip()
        gender = row.cells[gender_idx].text.strip()
        phone_number = row.cells[phone_number_idx].text.strip()
        month_born = row.cells[month_born_idx].text.strip()
        age = row.cells[age_idx].text.strip()
        occupation = row.cells[occupation_idx].text.strip()
        department_name = row.cells[department_name_idx].text.strip()
        member_status = row.cells[member_status_idx].text.strip()
        data_list.append({"firstName": first_name, "middleName": middle_name, "lastName": last_name, "gender": gender,"phoneNumber": phone_number, "monthBorn": month_born,"age": age, "occupation": occupation,"departmentName": department_name, "memberStatus": member_status})

    return data_list


@router.post("/helperFunction/1/add_new_members")
async def add_new_members(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        return {"error": "Invalid file type. Please upload a .docx file"}

    contents = await file.read()
    doc = Document(io.BytesIO(contents))
    data = read_docx1(doc)
    print("Extracted data:", data)

    # if "NAME" not in data[0] or "DEPARTMENT" not in data[0]:
    #     return {"error": "Document must contain 'Name:' and 'Department:' fields"}

    # result = update_department_in_db(data["NAME"], data["DEPARTMENT"])

    for row in data:
        print("rowss : ", row )
        # name = row.get("NAME")
        # department = row.get("DEPARTMENT")

        # if not name or not department:
        #     results.append({"name": name, "status": "error", "message": "Missing name or department"})
        #     continue

        update_result = add_member_in_db(row.get("firstName"), row.get("middleName"), row.get("lastName"),row.get("gender"),row.get("phoneNumber"),row.get("monthBorn"),row.get("age"),row.get("occupation"),row.get("departmentName"), row.get("memberStatus"))
        time.sleep(20) 

    # Return the appropriate message based on status
        if update_result["status"] == "added":
            print("message:", update_result["message"], "name:", row.get("firstName"))
        else:
            print("error:", update_result["message"], "name:", row.get("firstName"))



# gfhfhjgljkhlkk;kl;lk
def add_first_timer_in_db(name,phoneNumber, birthMonth, purposeOfComing):
    """Updates department for a given employee name in the database."""
    conn = get_connection()
    if not conn:
        return {"status": "error", "message": "Failed to connect to database"}

    cursor = conn.cursor()

    try:
        # Normalize name (lowercase + remove spaces)
        # clean_name = name.replace(" ", "").lower().strip()

        # Find matching members
        # cursor.execute("""
        #     SELECT id, "firstName", "middleName", "lastName"
        #     FROM members
        #     WHERE LOWER(TRIM("lastName" || COALESCE("middleName", '') || "firstName")) = %s
        # """, (clean_name,))

        # records = cursor.fetchall()

        # if len(records) == 0:
        #     msg = f"No record found for {name}"
        #     return {"status": "none", "message": msg}

        # elif len(records) > 1:
        #     msg = f"Multiple records found for {name}, skipping update."
        #     return {"status": "multiple", "message": msg}

        # Exactly one match → update department
        cursor.execute("SELECT COUNT(id) FROM members")
        count = cursor.fetchone()[0]

        ids = str(uuid.uuid4())

        cursor.execute(
            """INSERT INTO "FirstTimers" ("id","name", "phoneNumber", "birthMonth", "purposeOfComing") VALUES (%s, %s, %s, %s, %s)""",
            (ids, name, phoneNumber, birthMonth, purposeOfComing )
            )

        conn.commit()

        msg = f"Member updated for {name}"
        return {"status": "added", "message": msg}

    except OperationalError as e:
        msg = f"Database operation failed for {name}: {e}"
        print("⚠️", msg)
        return {"status": "error", "message": msg}

    finally:
        cursor.close()
        conn.close()


def read_docx2(doc):
    """Extracts name and department from a Word Document table."""
    data_list = []  # list of dicts for multiple rows

    # Check if the document has tables
    if not doc.tables:
        return {"error": "No tables found in document"}

    table = doc.tables[0]  # assuming first table contains your data

    # Get headers from first row
    headers = [cell.text.strip().upper() for cell in table.rows[0].cells]
    print(headers)

    # Check required columns
    # if "NAME" not in headers or "DEPARTMENT" not in headers:
    #     return {"error": "Table must contain 'NAME' and 'DEPARTMENT' columns"}

    name_idx = headers.index("NAME")
    phone_number_idx = headers.index("PHONENUMBER")
    birth_month_idx = headers.index("BIRTHMONTH")
    purpose_of_coming_idx = headers.index("PURPOSEOFCOMING")
    

    # Read remaining rows
    for row in table.rows[1:]:
        name  = row.cells[name_idx].text.strip()
        phone_number = row.cells[phone_number_idx].text.strip()
        birth_month = row.cells[birth_month_idx].text.strip()
        purpose_of_coming = row.cells[purpose_of_coming_idx].text.strip()

        data_list.append({"name": name, "phoneNumber": phone_number, "birthMonth": birth_month, "purposeOfComing": purpose_of_coming})

    return data_list


@router.post("/helperFunction/add_new_first_timers")
async def add_new_first_timers(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        return {"error": "Invalid file type. Please upload a .docx file"}

    contents = await file.read()
    doc = Document(io.BytesIO(contents))
    data = read_docx2(doc)
    print("Extracted data:", data)

    # if "NAME" not in data[0] or "DEPARTMENT" not in data[0]:
    #     return {"error": "Document must contain 'Name:' and 'Department:' fields"}

    # result = update_department_in_db(data["NAME"], data["DEPARTMENT"])

    for row in data:
        print("rowss : ", row )
        # name = row.get("NAME")
        # department = row.get("DEPARTMENT")

        # if not name or not department:
        #     results.append({"name": name, "status": "error", "message": "Missing name or department"})
        #     continue

        update_result = add_first_timer_in_db(row.get("name"), row.get("phoneNumber"), row.get("birthMonth"),row.get("purposeOfComing"))
        time.sleep(20) 

    # Return the appropriate message based on status
        if update_result["status"] == "added":
            print("message:", update_result["message"], "name:", row.get("name"))
        else:
            print("error:", update_result["message"], "name:", row.get("name"))


